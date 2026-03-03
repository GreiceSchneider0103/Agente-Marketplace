# marketplace_specialist/docx_builder.py
from docx import Document
from docx.shared import Pt
from datetime import datetime

def build_docx(payload: dict, product_data: dict) -> bytes:
    """
    Gera o DOCX final (neutro, sem logo).
    Retorna bytes do arquivo.
    """
    doc = Document()

    # Estilo base simples
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Anúncio Completo (Gerado pelo Agente)", level=1)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # --- ETAPA 1 (tabela) ---
    doc.add_heading("Etapa 1 — Dados do Produto", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    def add_row(k, v):
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = "" if v is None else str(v)

    # Campos do formulário (product_data)
    add_row("Nome do produto", product_data.get("nome_produto"))
    add_row("Marca / Linha", product_data.get("marca_linha"))
    add_row("Materiais", product_data.get("materiais"))
    add_row("Dimensões (LxAxP)", product_data.get("dimensoes"))
    add_row("Peso suportado (kg)", product_data.get("peso_suportado"))
    add_row("Cores disponíveis", product_data.get("cores_disponiveis"))
    add_row("Conteúdo da embalagem", product_data.get("conteudo_embalagem"))
    add_row("Necessita montagem?", product_data.get("necessita_montagem"))
    add_row("Tempo montagem", product_data.get("tempo_montagem"))
    add_row("Nível montagem", product_data.get("nivel_montagem"))
    add_row("Garantia", product_data.get("garantia"))
    add_row("Empresa", product_data.get("vendedor_empresa"))
    add_row("Canal / Marketplace", product_data.get("marketplace_alvo"))

    doc.add_paragraph("")

    # --- ETAPA 2 (estratégia) ---
    doc.add_heading("Etapa 2 — Análise Estratégica", level=2)
    ae = payload.get("analise_estrategica", {}) if isinstance(payload, dict) else {}

    doc.add_paragraph("Persona:")
    doc.add_paragraph(ae.get("persona", ""), style="List Bullet")

    doc.add_paragraph("Dores (3):")
    for d in (ae.get("dores") or []):
        doc.add_paragraph(str(d), style="List Bullet")

    doc.add_paragraph("Ganhos (3):")
    for g in (ae.get("ganhos") or []):
        doc.add_paragraph(str(g), style="List Bullet")

    doc.add_paragraph("Jornada de compra:")
    doc.add_paragraph(ae.get("jornada_compra", ""))

    doc.add_paragraph("Gatilhos mentais (3):")
    for gm in (ae.get("gatilhos_mentais") or []):
        doc.add_paragraph(str(gm), style="List Bullet")

    doc.add_paragraph("JTBD:")
    doc.add_paragraph(ae.get("jtbd", ""))

    doc.add_paragraph("PUV:")
    doc.add_paragraph(ae.get("puv", ""))

    doc.add_paragraph("Funcionalidades-chave (3–5):")
    for fc in (ae.get("funcionalidades_chave") or []):
        doc.add_paragraph(str(fc), style="List Bullet")

    doc.add_paragraph("Diferencial competitivo:")
    doc.add_paragraph(ae.get("diferencial_competitivo", ""))

    doc.add_paragraph("Prova social:")
    doc.add_paragraph(ae.get("prova_social", ""))

    # --- ETAPA 3 (SEO) ---
    doc.add_heading("Etapa 3 — SEO", level=2)
    seo = payload.get("seo", {}) if isinstance(payload, dict) else {}
    prim = seo.get("primarias") or []
    sec = seo.get("secundarias") or []
    tec = seo.get("termos_tecnicos") or []

    doc.add_paragraph("Palavras-chave primárias:")
    for x in prim:
        doc.add_paragraph(str(x), style="List Bullet")

    doc.add_paragraph("Palavras-chave secundárias:")
    for x in sec:
        doc.add_paragraph(str(x), style="List Bullet")

    doc.add_paragraph("Termos técnicos:")
    for x in tec:
        doc.add_paragraph(str(x), style="List Bullet")

    # --- Títulos / Modelo / Descrição ---
    doc.add_heading("Títulos (3)", level=2)
    for t in (payload.get("titulos") or []):
        doc.add_paragraph(str(t), style="List Number")

    doc.add_heading("Modelo", level=2)
    doc.add_paragraph(str(payload.get("modelo", ""))[:100])  # garante max 100

    doc.add_heading("Descrição completa", level=2)
    doc.add_paragraph(payload.get("descricao", ""))

    # --- Roteiro 7 imagens ---
    doc.add_heading("Roteiro — 7 Imagens", level=2)
    roteiro = payload.get("roteiro_imagens") or []
    for item in roteiro:
        if not isinstance(item, dict):
            continue
        numero = item.get("imagem") or item.get("numero")
        titulo = item.get("titulo", "")
        objetivo = item.get("objetivo", "")
        orient = item.get("orientacao_visual", "")
        texto = item.get("texto_sugerido") or item.get("texto") or ""

        doc.add_paragraph(f"Imagem {numero} — {titulo}".strip(), style="List Number")
        if objetivo:
            doc.add_paragraph(f"Objetivo: {objetivo}")
        if orient:
            doc.add_paragraph(f"Orientação visual: {orient}")
        if texto:
            doc.add_paragraph(f"Texto sugerido: {texto}")

    # Fontes (se você colocar no payload)
    doc.add_heading("Fontes e Referências", level=2)
    fontes = payload.get("fontes") or []
    if not fontes:
        doc.add_paragraph("Não informado.")
    else:
        for f in fontes:
            doc.add_paragraph(str(f), style="List Bullet")

    # exportar bytes
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()